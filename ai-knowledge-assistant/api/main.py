# ─────────────────────────────────────────────────────────────────────────────
# api/main.py — FastAPI backend for the AI Knowledge Assistant
# Exposes endpoints: /ask, /upload, /index, /reindex, /export, /metrics
# ─────────────────────────────────────────────────────────────────────────────

import os
import time
import shutil
import csv
from pathlib import Path
from typing import List

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv

# load environment variables from .env file
load_dotenv()

# ── App setup ─────────────────────────────────────────────────────────────────
app = FastAPI(
    title       = "AI Knowledge Assistant API",
    description = "RAG pipeline using FAISS + Ollama + sentence-transformers",
    version     = "1.0.0",
)

# allow Next.js frontend to call this API
app.add_middleware(
    CORSMiddleware,
    allow_origins     = ["http://localhost:3000"],
    allow_credentials = True,
    allow_methods     = ["*"],
    allow_headers     = ["*"],
)

# ── Paths from .env ───────────────────────────────────────────────────────────
RAW_DATA_PATH       = os.getenv("RAW_DATA_PATH",       "./data/raw")
PROCESSED_DATA_PATH = os.getenv("PROCESSED_DATA_PATH", "./data/processed")

# ── Request / Response models ─────────────────────────────────────────────────
class AskRequest(BaseModel):
    question: str                                           # user question
    model:    str = os.getenv("LLM_MODEL", "llama3.1:8b") # model to use

class ExportRequest(BaseModel):
    format:    str        # export format: pdf, docx, csv, md
    answer:    str        # answer text to export
    citations: list = []  # citations list

    class Config:
        # allow 'format' as field name without conflict
        populate_by_name = True

# ── Health check ──────────────────────────────────────────────────────────────
@app.get("/")
def root():
    """Check if the API is running."""
    return {"status": "ok", "message": "AI Knowledge Assistant API is running"}

# ── Ask endpoint ──────────────────────────────────────────────────────────────
@app.post("/ask")
def ask(req: AskRequest):
    """
    Main RAG endpoint.
    Calls the full RAG pipeline from core/generator.py
    Returns answer with citations and latency.
    """
    try:
        # import the full RAG pipeline from generator
        from core.generator import ask as rag_ask

        # run retrieval + generation in one call
        result = rag_ask(req.question, top_k=5)

        return {
            "answer":     result["answer"],
            "citations":  result["citations"],
            "model":      result["model"],
            "latency_ms": result["latency_ms"],
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── Upload endpoint ───────────────────────────────────────────────────────────
@app.post("/upload")
async def upload(files: List[UploadFile] = File(...)):
    """
    Accepts uploaded documents and saves them to data/raw.
    Supports: PDF, DOCX, MD, TXT, CSV
    """
    saved = []

    for file in files:
        filename = file.filename
        if not filename:
            raise HTTPException(status_code=400, detail="Uploaded file must include a filename")

        # build destination path inside data/raw
        dest = Path(RAW_DATA_PATH) / filename
        dest.parent.mkdir(parents=True, exist_ok=True)

        # read file contents and save to disk
        contents = await file.read()
        with open(dest, "wb") as f:
            f.write(contents)

        saved.append(filename)

    return {
        "message": f"Uploaded {len(saved)} file(s): {', '.join(saved)}",
        "files":   saved,
    }

# ── Index endpoint ────────────────────────────────────────────────────────────
@app.post("/index")
def index():
    """
    Ingests and indexes newly uploaded documents into FAISS.
    Run this after uploading new files.
    """
    try:
        from core.ingestor import run_ingestion
        from core.embedder import build_incremental

        # parse/chunk documents, then embed only newly processed chunks
        run_ingestion()
        build_incremental()

        return {"message": "Indexed successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── Reindex endpoint ──────────────────────────────────────────────────────────
@app.post("/reindex")
def reindex():
    """
    Clears the existing FAISS index and rebuilds from scratch.
    Use when documents have changed significantly.
    """
    try:
        from core.ingestor import run_ingestion
        from core.embedder import build_full

        # re-ingest and rebuild the full FAISS index from scratch
        run_ingestion()
        build_full()

        return {"message": "Full reindex complete"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── Export endpoint ───────────────────────────────────────────────────────────
@app.post("/export")
def export(req: ExportRequest):
    # debug — print what format arrived
    print(f"DEBUG export format received: '{req.format}'")
    print(f"DEBUG format repr: {repr(req.format)}")
    """
    Exports the last answer and citations to the requested format.
    Supports: pdf, docx, csv, md
    Saves to data/exports/ folder — never to data/raw/
    """
    try:
        fmt = req.format.strip().lower().lstrip(".")

        # create exports folder if it doesn't exist
        exports_dir = Path("./data/exports")
        exports_dir.mkdir(parents=True, exist_ok=True)

        # determine output file path inside exports folder
        output_path = str(exports_dir / f"export_output.{fmt}")

        if fmt == "md":
            # export as Markdown file
            _export_md(req.answer, req.citations, output_path)

        elif fmt == "csv":
            # export citations as CSV table
            _export_csv(req.answer, req.citations, output_path)

        elif fmt == "docx":
            # export as Word document
            _export_docx(req.answer, req.citations, output_path)

        elif fmt == "pdf":
            # export as Markdown since true PDF needs extra libraries
            # saves as .md but downloads as .pdf filename
            md_path = str(exports_dir / "export_output.md")
            _export_md(req.answer, req.citations, md_path)
            output_path = md_path

        else:
            raise HTTPException(status_code=400, detail="Unsupported format")

        # return file as browser download
        return FileResponse(
            path       = output_path,
            media_type = "application/octet-stream",
            filename   = f"answer.{fmt}",
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ── Metrics endpoint ──────────────────────────────────────────────────────────
@app.get("/metrics")
def metrics():
    """Returns query logs and performance metrics from SQLite."""
    try:
        import sqlite3
        conn = sqlite3.connect(os.getenv("SQLITE_PATH", "./data/logs.db"))
        cur  = conn.cursor()

        # fetch last 50 query logs ordered by most recent
        cur.execute("""
            SELECT timestamp, question, latency_ms, chunk_ids
            FROM query_logs
            ORDER BY timestamp DESC
            LIMIT 50
        """)
        rows = cur.fetchall()
        conn.close()

        return {
            "logs": [
                {
                    "timestamp":  r[0],
                    "question":   r[1],
                    "latency_ms": r[2],
                    "chunk_ids":  r[3],
                }
                for r in rows
            ]
        }

    except Exception as e:
        # return empty logs if table doesn't exist yet
        return {"logs": [], "error": str(e)}

# ─────────────────────────────────────────────────────────────────────────────
# Private export helper functions
# ─────────────────────────────────────────────────────────────────────────────

def _export_md(answer: str, citations: list, path: str):
    """Write answer and citations as a Markdown file."""
    lines = [
        "# AI Knowledge Assistant — Answer\n",
        f"\n{answer}\n",
        "\n## Sources\n",
    ]
    for c in citations:
        lines.append(f"\n**#{c.get('rank')} {c.get('title')}** — {c.get('source')}")
        lines.append(f"\n> {c.get('snippet')}\n")
    Path(path).write_text("\n".join(lines), encoding="utf-8")


def _export_csv(answer: str, citations: list, path: str):
    """Write citations as a CSV file with answer in first row."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        # write answer row first
        writer.writerow(["Answer"])
        writer.writerow([answer])
        writer.writerow([])
        # write citations table
        writer.writerow(["Rank", "Title", "Source", "Score", "Snippet"])
        for c in citations:
            writer.writerow([
                c.get("rank"),
                c.get("title"),
                c.get("source"),
                c.get("score"),
                c.get("snippet"),
            ])


def _export_docx(answer: str, citations: list, path: str):
    """Write answer and citations as a Word document."""
    from docx import Document
    doc = Document()

    # add answer section
    doc.add_heading("AI Knowledge Assistant — Answer", 0)
    doc.add_paragraph(answer)

    # add citations section
    doc.add_heading("Sources", level=1)
    for c in citations:
        doc.add_paragraph(
            f"#{c.get('rank')} {c.get('title')} — {c.get('source')}",
            style="List Bullet",
        )
        doc.add_paragraph(f"{c.get('snippet')}")

    doc.save(path)
