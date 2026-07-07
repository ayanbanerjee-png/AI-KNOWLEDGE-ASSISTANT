"""
core/generator.py — Task E: RAG Answer Generation (v2)
"""

from core.security import verify_token, redact_pii
from fastapi.responses import StreamingResponse
import io
import os
import time
import shutil
import argparse
from matplotlib import text
from pyparsing import line
import requests
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Depends
from pydantic import BaseModel

def _find_project_root() -> Path:
    here = Path(__file__).resolve()
    for parent in [here.parent, here.parent.parent, Path.cwd()]:
        if (parent / "data").exists():
            return parent
    return Path.cwd()

PROJECT_ROOT = _find_project_root()
load_dotenv(PROJECT_ROOT / ".env")

try:
    from core.retriever import retrieve, format_context
except ModuleNotFoundError:
    from retriever import retrieve, format_context

app = FastAPI(title="AI Knowledge Assistant", version="2.0.0")

OLLAMA_URL   = os.getenv("OLLAMA_URL",   "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "mistral")
TOP_K        = int(os.getenv("TOP_K",    "5"))
SNIPPET_LEN  = 200


def confidence_label(score: float) -> str:
    if score >= 0.65:   return "High"
    elif score >= 0.55: return "Medium"
    else:               return "Low"


def make_citations(results: list[dict]) -> list[dict]:
    citations = []
    for r in results:
        snippet = r["text"].strip().replace("\n", " ")
        if len(snippet) > SNIPPET_LEN:
            snippet = snippet[:SNIPPET_LEN].rsplit(" ", 1)[0] + " ..."
        citations.append({
            "rank":       r["rank"],
            "score":      r["score"],
            "confidence": confidence_label(r["score"]),
            "title":      r["title"],
            "source":     r["source"],
            "snippet":    snippet,
        })
    return citations


def format_citations_for_display(citations: list[dict]) -> str:
    lines = []
    for c in citations:
        lines.append(
            f"  [{c['rank']}] {c['title']}\n"
            f"       File      : {c['source']}\n"
            f"       Relevance : {c['confidence']} ({c['score']})\n"
            f"       Snippet   : \"{c['snippet']}\""
        )
    return "\n\n".join(lines)


SYSTEM_PROMPT = """You are a helpful AI Knowledge Assistant for an engineering team.
You answer questions strictly based on the provided context documents.
If the answer is not in the context, say "I don't have enough information to answer that."
Always be concise, accurate, and cite your sources by their Source number."""

DEVELOPER_PROMPT = """Use ONLY the information from the context below to answer the question.
Inline cite sources using [Source N] where relevant.
At the end, list used sources under a "Sources:" section like:
  - [Source N] Title (filename)"""


def build_prompt(question: str, context: str) -> str:
    return f"""{DEVELOPER_PROMPT}

---CONTEXT START---
{context}
---CONTEXT END---

Question: {question}

Answer:"""


def call_ollama(prompt: str, model: str = OLLAMA_MODEL) -> str:
    url     = f"{OLLAMA_URL}/api/generate"
    payload = {
        "model":  model,
        "prompt": prompt,
        "system": SYSTEM_PROMPT,
        "stream": False,
        "keep_alive": "10m",
        "options": {
            "temperature": 0.2,
            "top_p":       0.9,
            "num_predict": 512,
        }
    }
    try:
        r = requests.post(url, json=payload, timeout=180)
        r.raise_for_status()
        return r.json().get("response", "").strip()
    except requests.exceptions.ConnectionError:
        raise RuntimeError("Cannot connect to Ollama. Make sure it is running:\n  ollama serve")
    except requests.exceptions.Timeout:
        raise RuntimeError("Ollama timed out. The model may still be loading.")
    except Exception as e:
        raise RuntimeError(f"Ollama error: {e}")


def ask(question: str, top_k: int = TOP_K) -> dict:
    start   = time.time()
    results = retrieve(question, top_k=top_k)

    if not results:
        return {
            "question":   question,
            "answer":     "No relevant documents found in the knowledge base.",
            "citations":  [],
            "model":      OLLAMA_MODEL,
            "latency_ms": 0,
        }

    context    = format_context(results)
    prompt     = build_prompt(question, context)
    answer     = call_ollama(prompt)
    citations  = make_citations(results)
    latency_ms = int((time.time() - start) * 1000)


    # ── Log the query ──────────────────────────
    try:
        try:
            from core.logger import log_query
            from core.security import redact_pii
        except ModuleNotFoundError:
            from logger import log_query
            from security import redact_pii

        # Redact PII before logging
        safe_question, _ = redact_pii(question)
        
        log_query(
            question=safe_question,
            answer=answer,
            citations=citations,
            model=OLLAMA_MODEL,
            latency_ms=latency_ms,
            top_k=top_k,
            context=context,
        )
    except Exception:
        pass  # logging should never break the main flow


    return {
        "question":   question,
        "answer":     answer,
        "citations":  citations,
        "model":      OLLAMA_MODEL,
        "latency_ms": latency_ms,
    }


class AskRequest(BaseModel):
    question: str
    top_k:    int = TOP_K

class AskResponse(BaseModel):
    question:   str
    answer:     str
    citations:  list[dict]
    model:      str
    latency_ms: int

class ReindexRequest(BaseModel):
    mode: str = "update"


@app.get("/")
def root():
    return {"status": "ok", "message": "AI Knowledge Assistant is running"}


@app.post("/ask", response_model=AskResponse)
def ask_endpoint(req: AskRequest, _:bool = Depends(verify_token)):
    try:
        return ask(req.question, top_k=req.top_k)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal error: {e}")


@app.post("/upload")
async def upload_endpoint(file: UploadFile = File(...), _:bool = Depends(verify_token)):
    filename = file.filename
    if filename is None:
        raise HTTPException(status_code=400, detail="Uploaded file must include a filename")

    allowed = {".md", ".pdf", ".docx", ".csv"}
    ext     = Path(filename).suffix.lower()

    if ext not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    raw_dir  = PROJECT_ROOT / "data" / "raw"
    stem     = Path(filename).stem
    dest     = raw_dir / filename

    # ── Auto-rename if file already exists ──────────────
    counter = 1
    while dest.exists():
        new_name = f"{stem}({counter}){ext}"
        dest     = raw_dir / new_name
        counter += 1

    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    final_name = dest.name
    renamed    = final_name != filename

    return {
        "message": (
            f"✅ '{final_name}' uploaded successfully!"
            if not renamed else
            f"✅ Saved as '{final_name}' ('{filename}' already existed)"
        ),
        "skipped":  False,
        "filename": final_name,
        "renamed":  renamed,
    }

@app.post("/reindex")
def reindex_endpoint(req: ReindexRequest, _:bool = Depends(verify_token)):
    try:
        try:
            from core.ingestor import run_ingestion
            from core.embedder import build_incremental, build_full
        except ModuleNotFoundError:
            from ingestor import run_ingestion
            from embedder import build_incremental, build_full

        run_ingestion()
        if req.mode == "update":
            build_incremental()
        else:
            build_full()

        # ← add this
        try:
            from core.retriever import reload_vector_store
        except ModuleNotFoundError:
            from retriever import reload_vector_store
        reload_vector_store()

        return {"message": f"✅ Reindex ({req.mode}) complete"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Reindex failed: {str(e)}")


@app.post("/index")
async def index_endpoint(request: Request, _:bool = Depends(verify_token)):
    try:
        try:
            from core.ingestor import run_ingestion
            from core.embedder import build_incremental
        except ModuleNotFoundError:
            from ingestor import run_ingestion
            from embedder import build_incremental

        run_ingestion()
        build_incremental()

        try:
            from core.retriever import reload_vector_store
        except ModuleNotFoundError:
            from retriever import reload_vector_store
        reload_vector_store()

        return {"message": "✅ Indexing complete"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


"""
Export endpoint — add this to core/generator.py

Supports: docx, xlsx, pdf, json, md, txt
"""



@app.post("/export")
async def export_endpoint(request: Request, _:bool = Depends(verify_token)):
    """
    Export an answer with citations in the requested format.

    Request body:
    {
      "format":    "docx" | "xlsx" | "pdf" | "json" | "md" | "txt",
      "question":  "...",
      "answer":    "...",
      "citations": [...],
      "model":     "mistral",
      "latency_ms": 1234
    }
    """
    from fastapi.responses import StreamingResponse
    import io

    body      = await request.json()
    fmt       = str(body.get("format", "txt")).strip().lower().lstrip(".")
    question  = body.get("question", "")
    answer    = body.get("answer", "")
    citations = body.get("citations", [])
    model     = body.get("model", OLLAMA_MODEL)
    latency   = body.get("latency_ms", 0)

    if fmt == "json":
        import json as _json
        content = _json.dumps({
            "question":   question,
            "answer":     answer,
            "citations":  citations,
            "model":      model,
            "latency_ms": latency,
        }, indent=2, ensure_ascii=False)
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="application/json",
            headers={"Content-Disposition": "attachment; filename=answer.json"}
        )

    elif fmt == "txt" or fmt == "md":
        lines = [
            f"# Question\n{question}\n",
            f"# Answer\n{answer}\n",
            "# Sources\n",
        ]
        for c in citations:
            lines.append(f"- [{c.get('rank')}] {c.get('title')} ({c.get('source')}) — {c.get('confidence')} · {c.get('score')}")
            if c.get("snippet"):
                lines.append(f"  \"{c.get('snippet')}\"")
        lines.append(f"\n_Model: {model} | Latency: {latency}ms_")
        content = "\n".join(lines)
        ext = "md" if fmt == "md" else "txt"
        return StreamingResponse(
            io.BytesIO(content.encode("utf-8")),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=answer.{ext}"}
        )

    elif fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")

        doc = Document()

        # Title
        title = doc.add_heading("AI Knowledge Assistant — Export", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Question
        doc.add_heading("Question", level=1)
        doc.add_paragraph(question)

        # Answer
        doc.add_heading("Answer", level=1)
        doc.add_paragraph(answer)

        # Citations
        doc.add_heading("Sources", level=1)
        for c in citations:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{c.get('rank')}] {c.get('title')} ({c.get('source')})")
            run.bold = True
            p.add_run(f"\n  Relevance: {c.get('confidence')} · {c.get('score')}")
            if c.get("snippet"):
                p.add_run(f"\n  \"{c.get('snippet')}\"")

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Model: {model} | Latency: {latency}ms")
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=answer.docx"}
        )

    elif fmt == "csv":
        import csv

        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Question", question])
        writer.writerow(["Answer", answer])
        writer.writerow(["Model", model])
        writer.writerow(["Latency ms", latency])
        writer.writerow([])
        writer.writerow(["Rank", "Title", "Source", "Confidence", "Score", "Snippet"])

        for c in citations:
            writer.writerow([
                c.get("rank"),
                c.get("title"),
                c.get("source"),
                c.get("confidence"),
                c.get("score"),
                c.get("snippet", ""),
            ])

        return StreamingResponse(
            io.BytesIO(buf.getvalue().encode("utf-8-sig")),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=answer.csv"}
        )

    elif fmt == "xlsx":
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            raise HTTPException(status_code=500, detail="openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        assert ws is not None
        ws.title = "Answer"

        # Header style
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill("solid", fgColor="7C6AF7")

        # Question row
        ws.append(["Question", question])
        ws["A1"].font = Font(bold=True)

        # Answer row
        ws.append(["Answer", answer])
        ws["A2"].font = Font(bold=True)
        ws["B2"].alignment = Alignment(wrap_text=True)

        ws.append([])  # blank row

        # Citations header
        headers = ["Rank", "Title", "Source", "Confidence", "Score", "Snippet"]
        ws.append(headers)
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col)
            cell.font   = header_font
            cell.fill   = header_fill
            cell.alignment = Alignment(horizontal="center")

        # Citation rows
        for c in citations:
            ws.append([
                c.get("rank"),
                c.get("title"),
                c.get("source"),
                c.get("confidence"),
                c.get("score"),
                c.get("snippet", ""),
            ])

        # Column widths
        ws.column_dimensions["A"].width = 8
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 25
        ws.column_dimensions["D"].width = 12
        ws.column_dimensions["E"].width = 8
        ws.column_dimensions["F"].width = 50

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=answer.xlsx"}
        )

    elif fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        except ImportError:
            raise HTTPException(status_code=500, detail="reportlab not installed. Run: pip install reportlab")

        buf    = io.BytesIO()
        doc    = SimpleDocTemplate(buf, pagesize=A4,
                                   rightMargin=20*mm, leftMargin=20*mm,
                                   topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        story  = []

        accent = colors.HexColor("#7C6AF7")

        # Title
        title_style = ParagraphStyle("title", parent=styles["Title"],
                                     textColor=accent, fontSize=18)
        story.append(Paragraph("AI Knowledge Assistant", title_style))
        story.append(Spacer(1, 6*mm))

        # Question
        story.append(Paragraph("Question", styles["Heading1"]))
        story.append(Paragraph(question, styles["Normal"]))
        story.append(Spacer(1, 4*mm))

        # Answer
        story.append(Paragraph("Answer", styles["Heading1"]))

        def parse_answer_to_story(text: str, styles, accent, mm, Table, TableStyle, colors, Paragraph, Spacer):
            """Parse answer text — renders markdown tables as proper PDF tables."""
            lines      = text.split("\n")
            i          = 0
            while i < len(lines):
                line = lines[i]

                # Detect markdown table — starts with |
                if line.strip().startswith("|") and i + 1 < len(lines) and lines[i+1].strip().startswith("|---"):
                    # Collect all table rows
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith("|"):
                        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                        # Skip separator rows |---|---|
                        if not all(set(c.replace("-","").replace(":","").replace(" ","")) == set() for c in row):
                            table_lines.append(row)
                        i += 1

                    if table_lines:
                        # Build reportlab table
                        col_count  = max(len(r) for r in table_lines)
                        # Pad rows to same length
                        padded = [r + [""] * (col_count - len(r)) for r in table_lines]
                        col_w  = [160 / col_count * mm] * col_count
                        t = Table(padded, colWidths=col_w)
                        t.setStyle(TableStyle([
                            ("BACKGROUND", (0,0), (-1,0), accent),
                            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                            ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                            ("FONTNAME",   (0,1), (-1,-1), "Helvetica"),
                            ("FONTSIZE",   (0,0), (-1,-1), 8),
                            ("GRID",       (0,0), (-1,-1), 0.5, colors.grey),
                            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F5FF")]),
                            ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
                            ("PADDING",    (0,0), (-1,-1), 4),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 3*mm))
                else:
                    # Regular text line
                    if line.strip():
                        story.append(Paragraph(line.strip(), styles["Normal"]))
                    i += 1

        parse_answer_to_story(answer, styles, accent, mm, Table, TableStyle, colors, Paragraph, Spacer)
        story.append(Spacer(1, 4*mm))

        # Citations table
        if citations:
            story.append(Paragraph("Sources", styles["Heading1"]))
            table_data = [["#", "Title", "Source", "Confidence", "Score"]]
            for c in citations:
                table_data.append([
                    str(c.get("rank", "")),
                    c.get("title", "")[:30],
                    c.get("source", "")[:25],
                    c.get("confidence", ""),
                    str(c.get("score", "")),
                ])
            t = Table(table_data, colWidths=[10*mm, 55*mm, 50*mm, 25*mm, 20*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), accent),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",   (0,0), (-1,-1), 9),
                ("GRID",       (0,0), (-1,-1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F5FF")]),
                ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
            ]))
            story.append(t)

        # Footer
        story.append(Spacer(1, 6*mm))
        story.append(Paragraph(
            f"<font size=8 color='grey'>Model: {model} | Latency: {latency}ms</font>",
            styles["Normal"]
        ))

        doc.build(story)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=answer.pdf"}
        )

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")

@app.get("/metrics")
def metrics_endpoint():
    try:
        try:
            from core.logger import get_stats, get_recent_logs
        except ModuleNotFoundError:
            from logger import get_stats, get_recent_logs

        return {
            "stats":       get_stats(),
            "recent_logs": get_recent_logs(limit=10)
        }

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ask the AI Knowledge Assistant")
    parser.add_argument("question", type=str)
    parser.add_argument("--top-k",  type=int, default=TOP_K)
    parser.add_argument("--model",  type=str, default=OLLAMA_MODEL)
    args = parser.parse_args()

    print(f"\n💬  Question : {args.question}")
    print(f"    Model    : {args.model}")
    print(f"    Top-K    : {args.top_k}\n")
    print("⏳  Thinking ...\n")

    result = ask(args.question, top_k=args.top_k)

    print("=" * 60)
    print(f"✅  Answer  ({result['latency_ms']}ms)\n")
    print(result["answer"])
    print("\n--- Citations ---\n")
    print(format_citations_for_display(result["citations"]))
    print("=" * 60)
